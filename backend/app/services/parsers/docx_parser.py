"""
Word 文档 (DOCX) 解析器

使用 python-docx 进行 DOCX 解析
"""

import io
import logging
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from app.services.parsers.base import (
    BaseParser,
    DocumentMetadata,
    PageContent,
    ParsedDocument,
)

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Word 文档 (DOCX) 解析器"""

    supported_extensions: List[str] = [".docx"]
    supported_mime_types: List[str] = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ]

    async def parse(self, file_path: str) -> ParsedDocument:
        """
        解析 DOCX 文件

        Args:
            file_path: 文件路径

        Returns:
            解析后的文档对象
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install it with: pip install python-docx"
            )

        path = Path(file_path)
        stat = path.stat()

        doc = Document(file_path)

        return self._parse_docx_document(doc, path, stat.st_size)

    async def parse_bytes(self, data: bytes, filename: str) -> ParsedDocument:
        """
        从字节数据解析 DOCX 文件

        Args:
            data: 文件字节数据
            filename: 文件名

        Returns:
            解析后的文档对象
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install it with: pip install python-docx"
            )

        path = Path(filename)

        doc = Document(io.BytesIO(data))

        return self._parse_docx_document(doc, path, len(data))

    def _parse_docx_document(
        self,
        doc,
        path: Path,
        file_size: int,
    ) -> ParsedDocument:
        """
        解析 DOCX 文档对象

        Args:
            doc: python-docx Document 对象
            path: 文件路径
            file_size: 文件大小

        Returns:
            解析后的文档对象
        """
        # 提取核心属性
        core_props = doc.core_properties

        title = core_props.title or path.stem
        author = core_props.author
        created_at = core_props.created
        modified_at = core_props.modified

        # 提取段落内容
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格内容
        tables = []
        for table_idx, table in enumerate(doc.tables):
            table_content = self._extract_table_content(table)
            if table_content:
                tables.append(table_content)

        # 合并所有内容
        full_content = "\n\n".join(paragraphs)

        # 如果有表格，也加入内容
        if tables:
            full_content += "\n\n" + "\n\n".join(tables)

        word_count = self.count_words(full_content)
        language = self.detect_language(full_content)

        # 按章节分页（基于标题样式）
        pages = self._split_by_headings(doc)

        metadata = DocumentMetadata(
            title=title,
            author=author,
            created_at=created_at,
            modified_at=modified_at,
            file_type=".docx",
            file_size=file_size,
            word_count=word_count,
            language=language,
            custom_fields={
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "category": core_props.category,
                "comments": core_props.comments,
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
            },
        )

        return ParsedDocument(
            content=full_content,
            metadata=metadata,
            pages=pages,
        )

    def _extract_table_content(self, table) -> str:
        """
        提取表格内容为文本

        Args:
            table: python-docx Table 对象

        Returns:
            表格文本内容
        """
        rows_text = []

        for row in table.rows:
            cells_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells_text.append(cell_text)
            if cells_text:
                rows_text.append(" | ".join(cells_text))

        return "\n".join(rows_text)

    def _split_by_headings(self, doc) -> List[PageContent]:
        """
        按标题样式分割文档

        Args:
            doc: python-docx Document 对象

        Returns:
            页面内容列表
        """
        pages = []
        current_content = []
        page_number = 1

        heading_styles = ["Heading 1", "Heading 2", "标题 1", "标题 2"]

        for para in doc.paragraphs:
            # 检查是否是标题
            is_heading = False
            if para.style and para.style.name:
                is_heading = any(h in para.style.name for h in heading_styles)

            if is_heading and current_content:
                # 保存当前页面
                pages.append(
                    PageContent(
                        page_number=page_number,
                        content="\n\n".join(current_content),
                    )
                )
                page_number += 1
                current_content = []

            text = para.text.strip()
            if text:
                current_content.append(text)

        # 添加最后一部分
        if current_content:
            pages.append(
                PageContent(
                    page_number=page_number,
                    content="\n\n".join(current_content),
                )
            )

        # 如果没有按标题分割，将整个内容作为一页
        if not pages:
            all_text = "\n\n".join(
                para.text.strip() for para in doc.paragraphs if para.text.strip()
            )
            pages.append(
                PageContent(
                    page_number=1,
                    content=all_text,
                )
            )

        return pages
